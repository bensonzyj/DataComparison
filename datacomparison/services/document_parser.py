"""Document parsing utilities that convert files into raw text."""
from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Optional, Protocol

try:  # optional dependency for images
    from PIL import Image  # type: ignore
except Exception:  # pragma: no cover - dependency might be unavailable
    Image = None  # type: ignore

try:  # optional dependency for PDFs
    import pdfplumber  # type: ignore
except Exception:  # pragma: no cover - dependency might be unavailable
    pdfplumber = None

try:  # optional dependency for Word documents
    import docx  # type: ignore
except Exception:  # pragma: no cover
    docx = None

try:  # optional dependency for OCR
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover
    pytesseract = None


LOGGER = logging.getLogger(__name__)


class ParsedDocument(Dict[str, str]):
    """Simple dictionary subclass to store text representations."""


class Parser(Protocol):
    def parse(self, path: Path) -> ParsedDocument:
        ...


@dataclass
class ParserRegistry:
    """Registry for different document parsers."""

    parsers: Dict[str, Parser]

    def for_path(self, path: Path) -> Parser:
        suffix = path.suffix.lower()
        if suffix in self.parsers:
            return self.parsers[suffix]
        if suffix == ".jpeg":
            return self.parsers.get(".jpg")
        raise ValueError(f"No parser available for '{suffix}' files")


class TextParser:
    """Parser used for plain text files."""

    def parse(self, path: Path) -> ParsedDocument:
        text = path.read_text(encoding="utf-8")
        return ParsedDocument(text=text)


class PdfParser:
    """Parser for PDF files using pdfplumber."""

    def parse(self, path: Path) -> ParsedDocument:
        if pdfplumber is None:
            raise RuntimeError("pdfplumber is required for PDF parsing but is not installed")
        text_parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text_parts.append(page.extract_text() or "")
        return ParsedDocument(text="\n".join(text_parts))


class DocxParser:
    """Parser for Word documents using python-docx."""

    def parse(self, path: Path) -> ParsedDocument:
        if docx is None:
            raise RuntimeError("python-docx is required for DOCX parsing but is not installed")
        document = docx.Document(str(path))
        paragraphs = [para.text for para in document.paragraphs]
        return ParsedDocument(text="\n".join(paragraphs))


class ImageParser:
    """Parser for image files via OCR."""

    def parse(self, path: Path) -> ParsedDocument:
        if pytesseract is None:
            raise RuntimeError("pytesseract is required for OCR but is not installed")
        if Image is None:
            raise RuntimeError("Pillow is required for OCR but is not installed")
        image = Image.open(path)
        text = pytesseract.image_to_string(image)
        return ParsedDocument(text=text)


DEFAULT_REGISTRY = ParserRegistry(
    parsers={
        ".txt": TextParser(),
        ".pdf": PdfParser(),
        ".docx": DocxParser(),
        ".jpg": ImageParser(),
        ".png": ImageParser(),
    }
)


def parse_document(path: Path, registry: Optional[ParserRegistry] = None) -> ParsedDocument:
    registry = registry or DEFAULT_REGISTRY
    parser = registry.for_path(path)
    LOGGER.debug("Using parser %s for %s", parser.__class__.__name__, path)
    return parser.parse(path)
